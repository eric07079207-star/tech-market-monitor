from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
import shutil
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from src.project_memory import (  # noqa: E402
    ACTIVE_CONTEXT_FILE,
    CONVERSATION_LOG_FILE,
    DECISION_REGISTER_FILE,
    MEMORY_ARCHIVE_DIR,
    MEMORY_CHANGELOG_FILE,
    MEMORY_DIR,
    MEMORY_DOCX_FILE,
    PROJECT_MEMORY_FILE,
    load_memory_bundle,
)


DOC_FONT = "PingFangTC-Regular"


def set_run_font(run, name: str = DOC_FONT, size: float = 11, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def style_paragraph(paragraph, font_size: float = 11, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=font_size, bold=bold)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=False)
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)


def add_body_markdown(doc: Document, content: str) -> None:
    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(line[2:].strip())
            set_run_font(run)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.15
            continue
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        set_run_font(run)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.15


def add_dataframe_table(doc: Document, title: str, rows) -> None:
    add_heading(doc, title, level=2)
    if rows.empty:
        paragraph = doc.add_paragraph("目前沒有資料。")
        style_paragraph(paragraph)
        return
    table = doc.add_table(rows=1, cols=len(rows.columns), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, column in enumerate(rows.columns):
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(str(column))
        set_run_font(run, size=10, bold=True)
    for row in rows.itertuples(index=False):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run("" if value is None else str(value))
            set_run_font(run, size=10)


def archive_previous_docx(path: Path) -> None:
    if not path.exists():
        return
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    archive_path = MEMORY_ARCHIVE_DIR / f"專案記憶與討論摘要_{stamp}.docx"
    shutil.copy2(path, archive_path)


def build_doc() -> Path:
    bundle = load_memory_bundle()
    MEMORY_DIR.mkdir(parents=True, exist_ok=True)
    MEMORY_ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
    archive_previous_docx(MEMORY_DOCX_FILE)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title_run = title.add_run("專案記憶與討論摘要")
    set_run_font(title_run, size=20, bold=False)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(2)

    subtitle = doc.add_paragraph()
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    subtitle_run = subtitle.add_run(f"產生時間：{generated_at}｜來源：專案記憶資料夾")
    set_run_font(subtitle_run, size=10)
    subtitle.paragraph_format.space_after = Pt(10)

    intro = doc.add_paragraph(
        "這份文件整理目前專案的核心目標、重要討論、已確認決策、當前上下文與記憶更新紀錄。"
    )
    style_paragraph(intro)

    add_dataframe_table(doc, "專案記憶狀態", bundle.status_table)
    add_dataframe_table(doc, "最近更新", bundle.latest_updates)
    add_heading(doc, "長期記憶", level=1)
    add_body_markdown(doc, bundle.project_memory)
    add_heading(doc, "重要討論摘要", level=1)
    add_body_markdown(doc, bundle.conversation_log)
    add_heading(doc, "當前上下文", level=1)
    add_body_markdown(doc, bundle.active_context)

    if not bundle.decision_register.empty:
        add_dataframe_table(
            doc,
            "決策登錄",
            bundle.decision_register.rename(
                columns={
                    "decision_id": "編號",
                    "date": "日期",
                    "status": "狀態",
                    "category": "類別",
                    "title": "標題",
                    "decision": "決策",
                    "reason": "原因",
                    "impact_scope": "影響範圍",
                    "superseded_by": "被取代",
                }
            ),
        )

    if not bundle.memory_changelog.empty:
        add_dataframe_table(
            doc,
            "記憶更新紀錄",
            bundle.memory_changelog.rename(
                columns={
                    "date": "日期",
                    "change_summary": "更新內容",
                    "reason": "原因",
                    "source": "來源",
                }
            ),
        )

    source_note = doc.add_paragraph(
        f"來源檔案：{PROJECT_MEMORY_FILE.name}、{CONVERSATION_LOG_FILE.name}、{ACTIVE_CONTEXT_FILE.name}、"
        f"{DECISION_REGISTER_FILE.name}、{MEMORY_CHANGELOG_FILE.name}"
    )
    style_paragraph(source_note, font_size=9)

    doc.save(MEMORY_DOCX_FILE)
    return MEMORY_DOCX_FILE


if __name__ == "__main__":
    output = build_doc()
    print(output)
